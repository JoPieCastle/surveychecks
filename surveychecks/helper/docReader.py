import docx
from pathlib import Path


class docReader:
    def __init__(self, wordDocumentPath, infoIn="all"):
        self.wordDocument = Path(wordDocumentPath)
        self.text = ""
        self.readTableText()
        self.readBodyText()

    def readBodyText(self):
        """Get text from a word document main body

        Sets self.tableText:
            string: all text
        """
        doc = docx.Document(self.wordDocument)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)

        self.text += "\n".join(fullText)

    def readTableText(self):
        """Get text from a word document table body

        Uses:
            self.wordDocument (docx): Word document with variable ranges and filtercondition

        Sets self.tableText:
            string: all text from tables
        """
        doc = docx.Document(self.wordDocument)
        fullText = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    fullText.append(cell.text)
        self.text += "\n".join(fullText)

    def getText(self):
        return self.text
